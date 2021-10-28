import tkinter as tk
import tkinter.scrolledtext as S
from tkinter import messagebox as tkMessageBox
from tkinter import filedialog as tkFileDialog
from googletrans import Translator
from docx import Document
import os
import datetime
import requests

# ボタン1がクリックされた時の処理
def ButtonClick1():
    lines = input_box.get('1.0', 'end -1c')      # 入力欄に入力された文字列を取得

    # 改行前の「.」、「."」、「.”」、「：」をダミー文字に置き換え
    words = [".\n",".\"\n",".”\n",":\n"]

    for i, word in enumerate(words):
        lines = lines.replace(word,"XXX" + str(i))

    lines = lines.replace("-\n", "")            # 改行前の「-」を削除
    lines = lines.replace("\n", " ")            # 改行記号を削除

    # ダミー文字を元に戻してさらに空行を追加
    for i, word in enumerate(words):
        lines = lines.replace("XXX" + str(i), word + "\n")

    lines = lines.encode('utf-8', "ignore").decode('utf-8') # Pythonで取り扱えない文字を削除する。

    processed_box.delete('1.0', 'end')              # 整形結果欄をクリア
    processed_box.insert('1.0', lines)              # 整形結果欄に整形結果を出力

# ボタン2がクリックされた時の処理
def ButtonClick2():
    lines = get_text()

    while True:
        if len(lines) >= 5000:                  # 翻訳文字数が5000字以上の場合
            lines1 = lines[:5000].rsplit('\n\n', 1)[0]                  # 5000字以内の段落
            lines2 = lines[:5000].rsplit('\n\n', 1)[1] + lines[5000:]   # 残りの段落

            translator = Translator()
            lines1 = translator.translate(lines1, dest='ja').text       # Google翻訳
            translate_box.insert('end', lines1 + '\n\n')                # 翻訳結果欄に表示

            lines = lines2                      # 残りの段落を設定

        else:                                   # 翻訳文字数が5000字未満の場合
            translator = Translator()
            lines = translator.translate(lines, dest='ja').text         # Google翻訳
            translate_box.insert('end', lines)                          # 翻訳結果欄に表示

            break

# ボタン3がクリックされた時の処理
def ButtonClick3():
    lines = get_text()

    # ここはご自分で発行されたKEYを入れてください
    DEEPL_API_KEY = 'XXXXX'

    # URLクエリに仕込むパラメータの辞書を作っておく
    params = {
                "auth_key": DEEPL_API_KEY,
                "text": lines,
                "target_lang": 'JA'  # 出力テキストの言語を英語に設定
            }

    # パラメータと一緒にPOSTする
    request = requests.post("https://api.deepl.com/v2/translate", data=params)

    result = request.json()
    lines = result["translations"][0]["text"]
    # 翻訳結果欄に表示
    translate_box.insert('end', lines)

# ボタン4がクリックされた時の処理
def ButtonClick4():
    edit_text = processed_box.get('1.0', 'end -1c')                         # 形成結果欄に入力された文字列を取得

    translate_text = translate_box.get('1.0', 'end -1c')                # 翻訳結果欄に入力された文字列を取得

    fTyp=[('wordファイル',"*.docx")]                                        # Word対訳表テンプレートを選択
    iDir='.'
    filename=tkFileDialog.askopenfilename(filetypes=fTyp,initialdir=iDir)
    document = Document(filename)

    for paragraph in document.paragraphs:
        paragraph.text = paragraph.text.replace("原文をここに記載",edit_text)        # Word対訳表に原文を記載
        paragraph.text = paragraph.text.replace("訳文をここに記載",translate_text)   # Word対訳表に訳文を記載

    paragraphs = (paragraph
                  for table in document.tables
                  for row in table.rows
                  for cell in row.cells
                  for paragraph in cell.paragraphs)

    for paragraph in paragraphs:
        paragraph.text = paragraph.text.replace("原文をここに記載",edit_text)        # Word対訳表に原文を記載
        paragraph.text = paragraph.text.replace("訳文をここに記載",translate_text)   # Word対訳表に訳文を記載

    dt_now = datetime.datetime.now()                                    # Word対比表の保存
    dt_str = str(dt_now.hour).zfill(2)+str(dt_now.minute).zfill(2)+str(dt_now.second).zfill(2)
    savefilename = filename.replace(u".docx","_replace" + dt_str + ".docx")
    document.save(savefilename)

    # 保存結果を表示
    tkMessageBox.showinfo("作成完了",os.path.basename(savefilename) + " で保存しました。")

def get_text():
    # 整形結果欄に入力された文字列を取得
    lines = processed_box.get('1.0', 'end -1c')

    # 文字数を取得して画面に表示
    label = tk.Label(root, text = "文字数 "  + str(len(lines)) + "   ", font = ("Helvetica",14))
    label.place(relx = 0.58, y = 20)
    # 翻訳結果欄をクリア
    translate_box.delete('1.0', 'end')

    return lines

# メインのプログラム
root = tk.Tk()
root.geometry("1600x800")
root.title("Text_translator")

# ラベルの設定
label1 = tk.Label(root, text = "テキストを入力", font = ("Helvetica",14))
label1.place(x = 20, y = 20)
label2 = tk.Label(root, text = "形成結果", font = ("Helvetica",14))
label2.place(relx = 0.34, y = 20)
label3 = tk.Label(root, text = "翻訳結果", font = ("Helvetica",14))
label3.place(relx = 0.67, y = 20)

# ボタンの設定
button1 = tk.Button(root, text = "形成", font = ("Helvetica",14), command = ButtonClick1)
button1.place(x = 200, y = 15)
button2 = tk.Button(root, text = "Google翻訳", font = ("Helvetica",14), command = ButtonClick2)
button2.place(relx = 0.42, y = 15)
button3 = tk.Button(root, text = "DeepL翻訳", font = ("Helvetica",14), command = ButtonClick3)
button3.place(relx = 0.50, y = 15)
button4 = tk.Button(root, text = "Word対訳表作成", font = ("Helvetica",14), command = ButtonClick4)
button4.place(relx = 0.75, y = 15)

# 入力ボックスの設定
input_box = S.ScrolledText(root, font = ("Helvetica",12))
input_box.place(relheight = 0.89, relwidth = 0.32, relx = 0.01, y = 60)

# 整形結果ボックスの設定
processed_box = S.ScrolledText(root, font = ("Helvetica",12))
processed_box.place(relheight = 0.89, relwidth = 0.32, relx = 0.34, y = 60)

# 翻訳ボックスの設定
translate_box = S.ScrolledText(root, font = ("Helvetica",12))
translate_box.place(relheight = 0.89, relwidth = 0.32, relx = 0.67, y = 60)

root.mainloop()
